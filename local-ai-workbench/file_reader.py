"""Text extraction from uploaded student files.

Supported: .txt .md .pdf .docx .csv
We read each file and return its plain text so the LLM can grade it.

Every extractor is wrapped so a single bad file never kills the whole run.
"""
from __future__ import annotations

import csv
import io
from pathlib import Path

from config import config


class UnsupportedFileError(Exception):
    pass


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv"}
# Map file extension -> human-readable label for the UI
EXT_LABELS = {
    ".txt": "Plain text",
    ".md": "Markdown",
    ".pdf": "PDF",
    ".docx": "Word document",
    ".csv": "Spreadsheet (CSV)",
}


def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_csv(path: Path) -> str:
    with path.open(encoding="utf-8", errors="replace", newline="") as f:
        rows = list(csv.reader(f))
    lines = []
    for row in rows:
        lines.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(lines)


def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF. Uses pypdf if available; else reports clearly."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise UnsupportedFileError(
            "PDF support needs the 'pypdf' library. Install it with: "
            "pip install -r requirements.txt  (or pip install pypdf)"
        )
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    text = "\n".join(parts)
    if not text.strip():
        raise UnsupportedFileError(
            "This PDF appears to be a scanned image with no selectable text. "
            "PDFs that are just pictures cannot be read. Export the document "
            "as text, or use an OCR step first."
        )
    return text


def _extract_docx(path: Path) -> str:
    """Extract text from a .docx using python-docx."""
    try:
        import docx
    except ImportError:
        raise UnsupportedFileError(
            "Word (.docx) support needs the 'python-docx' library. Install it "
            "with: pip install -r requirements.txt"
        )
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    text = "\n".join(parts)
    if not text.strip():
        raise UnsupportedFileError(
            "This Word document appears to have no readable text."
        )
    return text


_EXTRACTORS = {
    ".txt": _extract_txt,
    ".md": _extract_txt,
    ".csv": _extract_csv,
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
}


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    """Extract text from raw uploaded bytes. Raises UnsupportedFileError."""
    path = Path(filename)
    ext = path.suffix.lower()
    if ext not in _EXTRACTORS:
        raise UnsupportedFileError(
            f"Unsupported file type '{ext}'. Upload one of: "
            + ", ".join(sorted(EXT_LABELS))
        )
    # Write to a temp file so the extractors keep a simple Path-based API.
    temp = Path(config.UPLOAD_DIR) / f"_tmp_{path.stem[:40]}{ext}"
    temp.write_bytes(data)
    try:
        return _EXTRACTORS[ext](temp)
    finally:
        try:
            temp.unlink()
        except Exception:
            pass
