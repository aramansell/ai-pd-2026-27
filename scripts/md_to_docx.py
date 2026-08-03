#!/usr/bin/env python3
"""Convert the AI-policy Markdown templates in downloads/ to clean .docx files.

Each .md follows a consistent structure:
  # Title
  *Italic note*
  **Department:** ...
  **School:** ...
  **Last updated:** ...
  **Focus:** ...
  ## Section
  paragraph text
  - bullet item
  - bullet item
"""
import os, re, glob
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "..", "downloads")   # markdown source
DST = SRC                                     # write docx alongside

BRAND = RGBColor(0x1D, 0x4E, 0xD8)   # deep blue, matches site
INK = RGBColor(0x1A, 0x23, 0x33)
MUTED = RGBColor(0x71, 0x80, 0x96)
TEAL = RGBColor(0x0F, 0x76, 0x6E)

FONT = "Calibri"


def set_base_styles(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(11)
    st.font.color.rgb = INK
    # ensure east-asian font mapping
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), FONT)


def add_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BRAND
    r.font.name = FONT
    p.space_after = Pt(6)
    return p


def add_subtitle(doc, text):
    # italic note line
    p = doc.add_paragraph()
    r = p.add_run(text.strip("*"))
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED
    p.space_after = Pt(12)
    return p


def add_meta(doc, line):
    # "**Label:** value" -> bold label, normal value
    m = re.match(r"\*\*(.+?)\*\*\s*:?\s*(.*)", line)
    p = doc.add_paragraph()
    if m:
        r1 = p.add_run(m.group(1).rstrip(":") + ":  ")
        r1.bold = True
        r1.font.color.rgb = INK
        p.add_run(m.group(2))
    else:
        p.add_run(line)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = TEAL
    r.font.name = FONT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    # thin bottom border under heading
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0F766E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_body(doc, text):
    # strip markdown emphasis markers for clean Word text
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    p = doc.add_paragraph()
    p.add_run(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, lines):
    for item in lines:
        item = item.strip().lstrip("-").strip()
        item = re.sub(r"\*\*(.+?)\*\*", r"\1", item)
        item = re.sub(r"\*(.+?)\*", r"\1", item)
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(3)


def convert(md_path):
    name = os.path.splitext(os.path.basename(md_path))[0]
    out_path = os.path.join(DST, name + ".docx")

    with open(md_path, encoding="utf-8") as f:
        raw = f.read()

    lines = raw.splitlines()
    doc = Document()
    set_base_styles(doc)

    bullets = []
    i = 0
    for ln in lines:
        s = ln.rstrip()
        if not s.strip():
            # flush any pending bullets
            if bullets:
                add_bullets(doc, bullets)
                bullets = []
            continue
        if s.startswith("# "):
            add_title(doc, s[2:].strip())
        elif s.startswith("## "):
            if bullets:
                add_bullets(doc, bullets)
                bullets = []
            add_heading(doc, s[3:].strip())
        elif s.startswith("*") and s.endswith("*") and len(s.strip("*").strip()) > 3:
            add_subtitle(doc, s)
        elif s.startswith("- "):
            bullets.append(s[2:])
        elif re.match(r"\*\*.+?\*\*", s):
            if bullets:
                add_bullets(doc, bullets)
                bullets = []
            add_meta(doc, s)
        else:
            if bullets:
                add_bullets(doc, bullets)
                bullets = []
            add_body(doc, s)
    if bullets:
        add_bullets(doc, bullets)

    doc.save(out_path)
    print(f"  ✓ {os.path.basename(out_path)}  ({os.path.getsize(out_path)} bytes)")
    return out_path


def main():
    mds = sorted(glob.glob(os.path.join(SRC, "policy-*.md"))) + \
          sorted(glob.glob(os.path.join(SRC, "classroom-agreement.md")))
    print("Converting markdown templates to .docx ...")
    for md in mds:
        convert(md)
    print("Done.")


if __name__ == "__main__":
    main()
